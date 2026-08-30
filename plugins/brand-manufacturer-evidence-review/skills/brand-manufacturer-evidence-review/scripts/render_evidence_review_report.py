from __future__ import annotations

import argparse
import json
import os
import re
import sys
from collections.abc import Iterable
from pathlib import Path
from typing import Any, Mapping, cast

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips

if __package__ in {None, ""}:
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from scripts.validate_evidence_review import assert_valid_payload
from scripts.check_output_filename import expected_names, validate as validate_output_filename


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_STYLE_PATH = ROOT / "assets" / "report-style" / "business-report-style.json"
REQUIRED_STYLE_TOKENS = {
    "page": {"size", "margin_top_mm", "margin_right_mm", "margin_bottom_mm", "margin_left_mm"},
    "fonts": {"east_asia", "latin"},
    "sizes_pt": {"title", "subtitle", "heading1", "heading2", "body", "table"},
    "spacing_pt": {"body_after", "heading1_before", "heading1_after", "heading2_before", "heading2_after"},
    "colors": {"primary", "secondary", "text", "muted", "border"},
    "table": {"cell_margin_twips", "header_repeat"},
}


def _is_number(value: Any) -> bool:
    return isinstance(value, (int, float)) and not isinstance(value, bool)


def _mapping(value: Any) -> Mapping[str, Any]:
    return value if isinstance(value, Mapping) else {}


def _mappings(value: Any) -> list[Mapping[str, Any]]:
    return [item for item in value if isinstance(item, Mapping)] if isinstance(value, list) else []


def _strings(value: Any) -> list[str]:
    return [item for item in value if isinstance(item, str)] if isinstance(value, list) else []


def _joined(values: Iterable[Any], separator: str = "；") -> str:
    return separator.join(str(value) for value in values if value is not None and str(value))


def _load_tokens(style_path: Path) -> Mapping[str, Any]:
    try:
        tokens = json.loads(style_path.read_text(encoding="utf-8"))
    except OSError as error:
        raise ValueError(f"STYLE_READ_ERROR {style_path}: {error}") from error
    except json.JSONDecodeError as error:
        raise ValueError(f"STYLE_JSON_ERROR {style_path}: {error}") from error
    if not isinstance(tokens, Mapping):
        raise ValueError("STYLE_TOKEN_INVALID style root must be an object")
    for group, keys in REQUIRED_STYLE_TOKENS.items():
        value = tokens.get(group)
        if not isinstance(value, Mapping):
            raise ValueError(f"STYLE_TOKEN_MISSING {group}")
        missing = sorted(keys - set(value))
        if missing:
            raise ValueError(f"STYLE_TOKEN_MISSING {group}.{','.join(missing)}")
    page = _mapping(tokens["page"])
    page_values = [page[key] for key in REQUIRED_STYLE_TOKENS["page"] - {"size"}]
    if page["size"] != "A4" or not all(_is_number(value) and 0 < value <= 100 for value in page_values):
        raise ValueError("STYLE_TOKEN_INVALID page")
    if page["margin_left_mm"] + page["margin_right_mm"] >= 210 or page["margin_top_mm"] + page["margin_bottom_mm"] >= 297:
        raise ValueError("STYLE_TOKEN_INVALID page margins exceed A4")
    fonts = _mapping(tokens["fonts"])
    sizes = _mapping(tokens["sizes_pt"])
    spacing = _mapping(tokens["spacing_pt"])
    colors = _mapping(tokens["colors"])
    table = _mapping(tokens["table"])
    if not all(isinstance(value, str) and value.strip() for value in fonts.values()):
        raise ValueError("STYLE_TOKEN_INVALID fonts")
    if not all(_is_number(value) and 0 < value <= 72 for value in sizes.values()):
        raise ValueError("STYLE_TOKEN_INVALID sizes_pt")
    if not all(_is_number(value) and 0 < value <= 72 for value in spacing.values()):
        raise ValueError("STYLE_TOKEN_INVALID spacing_pt")
    if not (_is_number(table["cell_margin_twips"]) and 0 < table["cell_margin_twips"] <= 1440):
        raise ValueError("STYLE_TOKEN_INVALID numeric tokens")
    if not isinstance(table["header_repeat"], bool):
        raise ValueError("STYLE_TOKEN_INVALID table.header_repeat")
    if not all(isinstance(value, str) and re.fullmatch(r"[0-9A-Fa-f]{6}", value) for value in colors.values()):
        raise ValueError("STYLE_TOKEN_INVALID colors")
    return tokens


def _set_run_font(run, tokens: Mapping[str, Any], size: float, color: str, *, bold: bool = False) -> None:
    fonts = _mapping(tokens["fonts"])
    run.font.name = str(fonts["latin"])
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), str(fonts["latin"]))
    r_fonts.set(qn("w:hAnsi"), str(fonts["latin"]))
    r_fonts.set(qn("w:eastAsia"), str(fonts["east_asia"]))
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _style_font(style, tokens: Mapping[str, Any], size: float, color: str, *, bold: bool = False) -> None:
    fonts = _mapping(tokens["fonts"])
    style.font.name = str(fonts["latin"])
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), str(fonts["latin"]))
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), str(fonts["latin"]))
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), str(fonts["east_asia"]))
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def apply_document_styles(document: Document, tokens: Mapping[str, Any]) -> None:
    page = _mapping(tokens["page"])
    sizes = _mapping(tokens["sizes_pt"])
    spacing = _mapping(tokens["spacing_pt"])
    colors = _mapping(tokens["colors"])
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(float(page["margin_top_mm"]))
    section.right_margin = Mm(float(page["margin_right_mm"]))
    section.bottom_margin = Mm(float(page["margin_bottom_mm"]))
    section.left_margin = Mm(float(page["margin_left_mm"]))
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal = document.styles["Normal"]
    _style_font(normal, tokens, float(sizes["body"]), str(colors["text"]))
    normal.paragraph_format.space_after = Pt(float(spacing["body_after"]))
    normal.paragraph_format.line_spacing = 1.2

    for name, size_key, before_key, after_key in (
        ("Heading 1", "heading1", "heading1_before", "heading1_after"),
        ("Heading 2", "heading2", "heading2_before", "heading2_after"),
    ):
        style = document.styles[name]
        _style_font(style, tokens, float(sizes[size_key]), str(colors["primary"]), bold=True)
        style.paragraph_format.space_before = Pt(float(spacing[before_key]))
        style.paragraph_format.space_after = Pt(float(spacing[after_key]))
        style.paragraph_format.keep_with_next = True

    subtitle = document.styles["Subtitle"]
    _style_font(subtitle, tokens, float(sizes["subtitle"]), str(colors["muted"]))
    subtitle.paragraph_format.space_after = Pt(float(spacing["body_after"]))
    subtitle.paragraph_format.line_spacing = 1.2
    document.settings.element.append(_update_fields_on_open())
    _configure_header_footer(section, tokens)


def _update_fields_on_open():
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    return update


def _configure_header_footer(section, tokens: Mapping[str, Any]) -> None:
    colors = _mapping(tokens["colors"])
    sizes = _mapping(tokens["sizes_pt"])
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("品牌-制造商证据复核报告")
    _set_run_font(run, tokens, float(sizes["subtitle"]), str(colors["muted"]))
    header.paragraph_format.space_after = Pt(0)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    _set_run_font(run, tokens, float(sizes["subtitle"]), str(colors["muted"]))
    _add_page_field(footer, tokens)
    run = footer.add_run(" 页")
    _set_run_font(run, tokens, float(sizes["subtitle"]), str(colors["muted"]))
    footer.paragraph_format.space_after = Pt(0)


def _add_page_field(paragraph, tokens: Mapping[str, Any]) -> None:
    run = paragraph.add_run()
    _set_run_font(run, tokens, float(_mapping(tokens["sizes_pt"])["subtitle"]), str(_mapping(tokens["colors"])["muted"]))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _add_text(document: Document, text: str, tokens: Mapping[str, Any], *, style: str = "Normal", bold: bool = False, color: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    sizes = _mapping(tokens["sizes_pt"])
    colors = _mapping(tokens["colors"])
    size = sizes["table"] if style == "Table Text" else sizes["body"]
    _set_run_font(run, tokens, float(size), color or str(colors["text"]), bold=bold)


def _add_label_value(document: Document, label: str, value: str, tokens: Mapping[str, Any]) -> None:
    paragraph = document.add_paragraph(style="Normal")
    colors = _mapping(tokens["colors"])
    sizes = _mapping(tokens["sizes_pt"])
    label_run = paragraph.add_run(f"{label}：")
    _set_run_font(label_run, tokens, float(sizes["body"]), str(colors["primary"]), bold=True)
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, tokens, float(sizes["body"]), str(colors["text"]))


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, margin: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "bottom", "start", "end"):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(margin))
        element.set(qn("w:type"), "dxa")


def _set_width(element, tag: str, width: int) -> None:
    child = element.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        element.append(child)
    child.set(qn("w:type"), "dxa")
    child.set(qn("w:w"), str(width))


def _apply_table_geometry(table, widths: list[int], tokens: Mapping[str, Any]) -> None:
    if any(width <= 0 for width in widths):
        raise ValueError("table column width must be positive")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    content_width = int(round((table._parent.part.document.sections[0].page_width.twips - table._parent.part.document.sections[0].left_margin.twips - table._parent.part.document.sections[0].right_margin.twips)))
    indent = int(_mapping(tokens["table"])["cell_margin_twips"])
    table_width = sum(widths)
    if indent + table_width > content_width:
        raise ValueError("table geometry exceeds available width")
    table_properties = table._tbl.tblPr
    _set_width(table_properties, "w:tblW", table_width)
    _set_width(table_properties, "w:tblInd", indent)
    layout = table_properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    margin = int(_mapping(tokens["table"])["cell_margin_twips"])
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            _set_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[index])
            _set_cell_margins(cell, margin)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row, tokens: Mapping[str, Any]) -> None:
    if not bool(_mapping(tokens["table"])["header_repeat"]):
        return
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_table_borders(table, color: str) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), color)
        borders.append(border)
    table._tbl.tblPr.append(borders)


def _add_table(document: Document, headers: list[str], rows: list[list[str]], weights: list[int], tokens: Mapping[str, Any]) -> None:
    content_width = int(round((document.sections[0].page_width.twips - document.sections[0].left_margin.twips - document.sections[0].right_margin.twips)))
    indent = int(_mapping(tokens["table"])["cell_margin_twips"])
    available_width = content_width - indent
    if available_width <= 0:
        raise ValueError("table geometry exceeds available width")
    widths = [int(available_width * weight / sum(weights)) for weight in weights]
    widths[-1] += available_width - sum(widths)
    table = document.add_table(rows=1, cols=len(headers))
    _set_table_borders(table, str(_mapping(tokens["colors"])["border"]))
    header_row = table.rows[0]
    for index, value in enumerate(headers):
        cell = header_row.cells[index]
        _set_cell_shading(cell, str(_mapping(tokens["colors"])["secondary"]))
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        _set_run_font(run, tokens, float(_mapping(tokens["sizes_pt"])["table"]), str(_mapping(tokens["colors"])["primary"]), bold=True)
    _repeat_header(header_row, tokens)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            _set_run_font(run, tokens, float(_mapping(tokens["sizes_pt"])["table"]), str(_mapping(tokens["colors"])["text"]))
    _apply_table_geometry(table, widths, tokens)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _entity_names(payload: Mapping[str, Any]) -> dict[str, str]:
    return {str(entity.get("主体ID")): str(entity.get("规范名称")) for entity in _mappings(payload.get("主体索引"))}


def _relationship_lines(review: Mapping[str, Any], entities: Mapping[str, str]) -> list[str]:
    lines: list[str] = []
    role_groups = (("主体关系", "商标权利人"), ("主体关系", "品牌运营主体"), ("主体关系", "母公司"), ("主体关系", "收购与历史关系"), ("制造关系", "品牌层面主要制造商"), ("制造关系", "具体SKU制造商"))
    for group, role in role_groups:
        for link in _mappings(_mapping(review.get(group)).get(role)):
            entity_id = str(link.get("主体ID"))
            details = [str(link.get("结论状态")), str(link.get("可靠性等级"))]
            evidence = _joined(_strings(link.get("证据引用")), "、")
            if evidence:
                details.append(f"证据：{evidence}")
            products = _joined(_strings(link.get("适用商品ID")), "、")
            if products:
                details.append(products)
            limitations = _joined(_strings(link.get("适用限制")))
            if limitations:
                details.append(limitations)
            lines.append(f"{role}：{entities.get(entity_id, entity_id)}（{'；'.join(details)}）")
    return lines


def add_cover(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    report = _mapping(payload.get("报告信息"))
    scope = _mapping(payload.get("调查范围"))
    colors = _mapping(tokens["colors"])
    sizes = _mapping(tokens["sizes_pt"])
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(56)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(str(report.get("报告标题")))
    _set_run_font(run, tokens, float(sizes["title"]), str(colors["primary"]), bold=True)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(str(scope.get("产品规范名称")))
    _add_label_value(document, "调查日期", str(report.get("调查日期")), tokens)
    _add_label_value(document, "底稿编号", str(report.get("底稿编号")), tokens)
    _add_label_value(document, "生成时间", str(report.get("生成时间")), tokens)
    document.add_page_break()


def add_execution_summary(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    scope = _mapping(payload.get("调查范围"))
    document.add_heading("执行摘要", level=1)
    _add_label_value(document, "产品原始名称", str(scope.get("产品原始名称")), tokens)
    _add_label_value(document, "产品规范名称", str(scope.get("产品规范名称")), tokens)
    _add_label_value(document, "产品描述", str(scope.get("产品描述")), tokens)
    for review in _mappings(payload.get("品牌复核结果")):
        brand = _mapping(review.get("品牌"))
        _add_label_value(document, str(brand.get("规范名称")), str(review.get("查询结果摘要")), tokens)


def add_evidence_chain(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    # Method and search logs remain in JSON for auditability; the DOCX is result-focused.
    return


def add_claim_matrix(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("主张与证据矩阵", level=1)
    rows = []
    for claim in _mappings(payload.get("主张证据矩阵")):
        rows.append([
            str(claim.get("主张编号")),
            str(claim.get("主张")),
            _joined(_strings(claim.get("证据引用")), "、"),
            str(claim.get("证据强度")),
            str(claim.get("结论状态")),
            str(claim.get("结论边界")),
            str(claim.get("下一步补证")),
            str(claim.get("冲突说明") or "无"),
        ])
    _add_table(document, ["主张编号", "主张", "证据", "强度", "状态", "结论边界", "下一步补证", "冲突"], rows, [9, 22, 9, 9, 9, 18, 14, 10], tokens)


def add_sku_checklists(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    products = _mappings(_mapping(payload.get("调查范围")).get("代表性商品"))
    if not products:
        return
    document.add_heading("SKU证据完整度", level=1)
    keys = ("包装标签", "产品铭牌", "型号或UPC", "说明书Manufacturer", "说明书Importer", "平台销售字段", "合规或监管文件")
    rows = []
    for product in products:
        checklist = _mapping(product.get("SKU证据核验"))
        rows.append([
            str(product.get("平台标识符") or product.get("商品ID")),
            str(checklist.get("证据完整度")),
            _joined((f"{key}:{_mapping(checklist.get(key)).get('状态')}" for key in keys), "；"),
        ])
    _add_table(document, ["SKU/平台标识", "完整度", "逐项状态"], rows, [22, 15, 63], tokens)


def add_brand_sections(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("各品牌逐一复核结果", level=1)
    target_order = [(str(brand.get("原始名称")), str(brand.get("规范名称"))) for brand in _mappings(_mapping(payload.get("调查范围")).get("目标品牌"))]
    reviews = {tuple(str(_mapping(review.get("品牌")).get(key)) for key in ("原始名称", "规范名称")): review for review in _mappings(payload.get("品牌复核结果"))}
    entities = _entity_names(payload)
    for brand_key in target_order:
        review = reviews[brand_key]
        document.add_heading(brand_key[1], level=2)
        _add_label_value(document, "查询结果摘要", str(review.get("查询结果摘要")), tokens)
        for line in _relationship_lines(review, entities):
            _add_text(document, line, tokens)
        manufacturing = _mapping(review.get("制造关系"))
        _add_label_value(document, "制造模式", str(manufacturing.get("制造模式")), tokens)
        for limitation in _strings(manufacturing.get("SKU适用限制")):
            _add_label_value(document, "SKU适用限制", limitation, tokens)
        assessment = _mapping(review.get("结论评价"))
        for label in ("可靠性等级", "可靠性依据", "人工复核建议"):
            _add_label_value(document, label, str(assessment.get(label)), tokens)
        for note in _strings(review.get("关键说明")):
            _add_label_value(document, "关键说明", note, tokens)
        source_rows = []
        for source in _mappings(review.get("主要来源")):
            source_rows.append([
                str(source.get("证据编号")),
                str(source.get("来源名称")),
                str(source.get("URL")),
                str(source.get("来源类别")),
                str(source.get("访问日期")),
                _joined(_strings(source.get("支持结论"))),
                str(source.get("证据等级")),
                str(source.get("原文摘录")),
                str(source.get("页面定位")),
            ])
        _add_table(document, ["证据编号", "来源名称", "URL", "来源类别", "访问日期", "支持结论", "证据等级", "原文摘录", "页面定位"], source_rows, [8, 13, 18, 11, 9, 13, 8, 14, 6], tokens)


def add_entity_summary(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("法律主体汇总", level=1)
    brands = {str(item.get("品牌ID")): str(item.get("规范名称")) for item in _mappings(_mapping(payload.get("调查范围")).get("目标品牌"))}
    rows = []
    for entity in sorted(_mappings(payload.get("主体索引")), key=lambda value: str(value.get("主体ID"))):
        rows.append([
            str(entity.get("主体ID")),
            str(entity.get("规范名称")),
            _joined((brands.get(identifier, identifier) for identifier in _strings(entity.get("关联品牌"))), "、"),
            _joined(_strings(entity.get("主体角色")), "、"),
            str(entity.get("当前状态")),
            str(entity.get("可靠性等级")),
            _joined(_strings(entity.get("适用限制"))),
        ])
    _add_table(document, ["主体ID", "规范名称", "关联品牌", "主体角色", "状态", "可靠性", "重要限制"], rows, [10, 19, 13, 16, 8, 8, 26], tokens)


def add_overall_classification(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("总体结论与人工复核分级", level=1)
    names = {str(item.get("品牌ID")): str(item.get("规范名称")) for item in _mappings(_mapping(payload.get("调查范围")).get("目标品牌"))}
    rows = [[str(item.get("复核等级")), _joined((names.get(identifier, identifier) for identifier in _strings(item.get("品牌"))), "、"), str(item.get("适用结论"))] for item in _mappings(payload.get("总体复核分级"))]
    _add_table(document, ["复核等级", "品牌", "适用结论"], rows, [15, 25, 60], tokens)


def add_database_access(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("数据库访问", level=1)
    rows = [[str(item.get(key)) for key in ("数据库", "官方入口", "访问时间", "访问结果", "使用建议")] for item in _mappings(payload.get("数据库访问记录"))]
    _add_table(document, ["数据库", "官方入口", "访问时间", "访问结果", "使用建议"], rows, [18, 25, 16, 21, 20], tokens)


def add_limitations(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("使用限制与后续建议", level=1)
    for limitation in _strings(payload.get("使用限制与建议")):
        _add_text(document, limitation, tokens)


def add_report_note(document: Document, payload: Mapping[str, Any], tokens: Mapping[str, Any]) -> None:
    document.add_heading("报告说明", level=1)
    _add_label_value(document, "Skill版本", str(payload.get("Skill版本")), tokens)
    _add_label_value(document, "规范版本", str(payload.get("规范版本")), tokens)
    confirmation = _mapping(payload.get("用户确认"))
    _add_label_value(document, "确认时间", str(confirmation.get("确认时间")), tokens)
    _add_label_value(document, "用户确认原文", str(confirmation.get("用户确认原文")), tokens)
    _add_label_value(document, "内容摘要哈希", str(confirmation.get("内容摘要哈希")), tokens)


def render_report(payload: Mapping[str, Any], output_path: Path, style_path: Path | None = None) -> Path:
    _, expected_docx = expected_names(cast(dict, payload))
    if output_path.name != expected_docx:
        raise ValueError(f"DOCX_FILENAME_MISMATCH expected={expected_docx} actual={output_path.name}")
    assert_valid_payload(payload, require_confirmed=True)
    output_path = Path(output_path)
    tokens = _load_tokens(Path(style_path) if style_path is not None else DEFAULT_STYLE_PATH)
    document = Document()
    apply_document_styles(document, tokens)
    add_cover(document, payload, tokens)
    add_execution_summary(document, payload, tokens)
    add_evidence_chain(document, payload, tokens)
    add_claim_matrix(document, payload, tokens)
    add_sku_checklists(document, payload, tokens)
    add_brand_sections(document, payload, tokens)
    add_entity_summary(document, payload, tokens)
    add_overall_classification(document, payload, tokens)
    add_database_access(document, payload, tokens)
    add_limitations(document, payload, tokens)
    add_report_note(document, payload, tokens)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = output_path.with_name(f".{output_path.name}.tmp")
    try:
        document.save(temporary_path)
        os.replace(temporary_path, output_path)
    finally:
        temporary_path.unlink(missing_ok=True)
    return output_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Render a confirmed evidence review JSON payload to DOCX.")
    parser.add_argument("json_path", type=Path)
    parser.add_argument("docx_path", type=Path)
    parser.add_argument("--style-path", type=Path)
    args = parser.parse_args(argv)
    try:
        payload = json.loads(args.json_path.read_text(encoding="utf-8"))
        validate_output_filename(payload, args.json_path, args.docx_path)
        render_report(cast(Mapping[str, Any], payload), args.docx_path, args.style_path)
    except (OSError, json.JSONDecodeError, ValueError) as error:
        print(error)
        return 1
    print(f"RENDERED: {args.docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
