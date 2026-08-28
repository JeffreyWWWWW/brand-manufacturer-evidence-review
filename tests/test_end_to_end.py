import re
import xml.etree.ElementTree as ET
from zipfile import ZipFile

from docx import Document

from scripts.render_evidence_review_report import render_report
from scripts.validate_evidence_review import compute_content_hash, validate_payload
from tests.helpers import cloned_fixture


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def confirmed_payload():
    payload = cloned_fixture("shared-entity-review.json")
    payload["用户确认"] = {
        "是否确认": True,
        "确认时间": "2026-08-28T10:00:00+08:00",
        "用户确认原文": "确认该合成案例中的品牌、主体关系、制造范围及已披露限制",
        "内容摘要哈希": None,
    }
    payload["用户确认"]["内容摘要哈希"] = compute_content_hash(payload)
    return payload


def document_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def test_synthetic_fixture_validates_and_renders_from_confirmed_json(tmp_path):
    payload = confirmed_payload()
    assert validate_payload(payload, require_confirmed=True) == []

    expected_brands = ["星驰（STARRY）", "远航（VOYAGE）"]
    assert [brand["规范名称"] for brand in payload["调查范围"]["目标品牌"]] == expected_brands
    shared_entities = [
        entity for entity in payload["主体索引"] if entity["关联品牌"] == ["BRD-001", "BRD-002"]
    ]
    assert len(shared_entities) == 1
    assert set(shared_entities[0]["主体角色"]) == {"商标权利人", "品牌运营主体"}

    output = render_report(payload, tmp_path / "brand-manufacturer-evidence-review.docx")
    document = Document(output)
    text = document_text(document)
    brand_headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 2" and paragraph.text in expected_brands
    ]
    assert brand_headings == expected_brands
    assert shared_entities[0]["规范名称"] in text
    assert "证据：EVD-001" in text
    assert "证据：EVD-002" in text
    assert all(
        source["URL"] in text
        for review in payload["品牌复核结果"]
        for source in review["主要来源"]
    )

    section = document.sections[0]
    assert "品牌-制造商证据复核报告" in "\n".join(
        paragraph.text for paragraph in section.header.paragraphs
    )
    assert "第 1 页" in "\n".join(paragraph.text for paragraph in section.footer.paragraphs)
    with ZipFile(output) as archive:
        footer = ET.fromstring(archive.read("word/footer1.xml"))
        root = ET.fromstring(archive.read("word/document.xml"))
    assert root.findall(".//w:pStyle[@w:val='Heading1']", NS)
    assert root.findall(".//w:pStyle[@w:val='Heading2']", NS)
    assert footer.findall(".//w:instrText[.=' PAGE ']", NS)
    assert not re.search(r"stage_|ready_for_|intake", text, flags=re.IGNORECASE)


def test_renderer_output_depends_only_on_payload_facts(tmp_path):
    payload = confirmed_payload()
    payload["品牌复核结果"][0]["关键说明"].append("仅由端到端 JSON 提供的哨兵事实")
    payload["用户确认"]["内容摘要哈希"] = compute_content_hash(payload)

    output = render_report(payload, tmp_path / "report.docx")

    assert "仅由端到端 JSON 提供的哨兵事实" in document_text(Document(output))


def test_confirmation_hash_detects_fact_changes():
    payload = confirmed_payload()
    original_hash = payload["用户确认"]["内容摘要哈希"]

    payload["报告信息"]["报告标题"] = "已变更事实"

    assert compute_content_hash(payload) != original_hash
    assert "CONFIRMATION_HASH_MISMATCH" in {
        issue.code for issue in validate_payload(payload, require_confirmed=True)
    }
