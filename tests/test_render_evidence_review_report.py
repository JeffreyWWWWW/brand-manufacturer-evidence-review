import json
import subprocess
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from scripts.render_evidence_review_report import DEFAULT_STYLE_PATH, render_report
from scripts.validate_evidence_review import compute_content_hash
from tests.helpers import SKILL_ROOT, cloned_fixture


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def document_text(document):
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + cells)


def confirmed_payload():
    payload = cloned_fixture("minimal-valid-review.json")
    payload["用户确认"] = {
        "是否确认": True,
        "确认时间": "2026-08-27T10:00:00+08:00",
        "用户确认原文": "确认以上品牌与主体关系",
        "内容摘要哈希": compute_content_hash(payload),
    }
    return payload


def xml_part(path, name):
    with ZipFile(path) as archive:
        return archive.read(name)


def replace_string_values(value, old, new):
    if isinstance(value, dict):
        return {key: replace_string_values(item, old, new) for key, item in value.items()}
    if isinstance(value, list):
        return [replace_string_values(item, old, new) for item in value]
    return new if value == old else value


def test_renderer_rejects_unconfirmed_payload_without_writing_output(tmp_path):
    output = tmp_path / "out.docx"
    with pytest.raises(ValueError, match="NOT_CONFIRMED"):
        render_report(cloned_fixture("minimal-valid-review.json"), output)
    assert not output.exists()


def test_renderer_writes_expected_sections_entity_summary_and_sources(tmp_path):
    payload = confirmed_payload()
    output = render_report(payload, tmp_path / "out.docx")
    document = Document(output)
    text = document_text(document)
    for heading in (
        "执行摘要",
        "证据链方法",
        "各品牌逐一复核结果",
        "法律主体汇总",
        "总体结论与人工复核分级",
        "数据库访问",
        "使用限制与后续建议",
        "报告说明",
    ):
        assert heading in text
    assert payload["主体索引"][0]["规范名称"] in text
    source = payload["品牌复核结果"][0]["主要来源"][0]
    for field in ("来源名称", "URL", "访问日期", "证据等级"):
        assert source[field] in text
    assert source["支持结论"][0] in text


def test_renderer_writes_all_source_fields_with_sentinels(tmp_path):
    payload = confirmed_payload()
    source = payload["品牌复核结果"][0]["主要来源"][0]
    payload = replace_string_values(payload, source["证据编号"], "EVD-999")
    source = payload["品牌复核结果"][0]["主要来源"][0]
    source["来源类别"] = "SENTINEL-SOURCE-CATEGORY"
    source["URL"] = "https://example.test/" + "long-path-segment-" * 20
    payload["用户确认"]["内容摘要哈希"] = compute_content_hash(payload)

    document = Document(render_report(payload, tmp_path / "out.docx"))
    text = document_text(document)
    assert "EVD-999" in text
    assert "SENTINEL-SOURCE-CATEGORY" in text
    source_table = document.tables[0]
    assert len(source_table.columns) == 7
    assert source["URL"] in text
    assert not source_table.rows[1].cells[2]._tc.find(qn("w:noWrap"))
    assert [cell.text for cell in source_table.rows[0].cells] == [
        "证据编号",
        "来源名称",
        "URL",
        "来源类别",
        "访问日期",
        "支持结论",
        "证据等级",
    ]


def test_renderer_preserves_target_brand_order_and_entity_id_order(tmp_path):
    payload = cloned_fixture("shared-entity-review.json")
    first_entity = payload["主体索引"][0]
    payload["品牌复核结果"][1]["主体关系"]["品牌运营主体"][0]["主体ID"] = "ENT-002"
    payload["主体关系"][1]["终点ID"] = "ENT-002"
    payload["主体索引"] = [
        {
            **first_entity,
            "主体ID": "ENT-002",
            "规范名称": "排序在后主体",
            "关联品牌": ["BRD-002"],
            "主体角色": ["品牌运营主体"],
            "证据引用": ["EVD-002"],
        },
        {
            **first_entity,
            "主体ID": "ENT-001",
            "关联品牌": ["BRD-001"],
            "主体角色": ["商标权利人"],
            "证据引用": ["EVD-001"],
        },
    ]
    for relation, review in zip(payload["主体关系"], payload["品牌复核结果"]):
        if relation["关系类型"] == "商标权利人":
            relation["适用限制"] = list(review["主体关系"]["商标权利人"][0]["适用限制"])
        else:
            relation["适用限制"] = list(review["主体关系"]["品牌运营主体"][0]["适用限制"])
    payload["用户确认"] = {
        "是否确认": True,
        "确认时间": "2026-08-27T10:00:00+08:00",
        "用户确认原文": "确认共享主体结果",
        "内容摘要哈希": compute_content_hash(payload),
    }
    document = Document(render_report(payload, tmp_path / "out.docx"))
    text = document_text(document)
    assert text.index("星驰（STARRY）") < text.index("远航（VOYAGE）")
    entity_rows = [row.cells[0].text for table in document.tables for row in table.rows if row.cells]
    assert entity_rows.index("ENT-001") < entity_rows.index("ENT-002")


def test_renderer_does_not_add_facts_absent_from_json(tmp_path):
    payload = confirmed_payload()
    payload["品牌复核结果"][0]["关键说明"] = ["仅来自结构化 JSON 的哨兵文本"]
    payload["用户确认"]["内容摘要哈希"] = compute_content_hash(payload)
    document = Document(render_report(payload, tmp_path / "out.docx"))
    all_text = document_text(document)
    assert "仅来自结构化 JSON 的哨兵文本" in all_text
    assert "待补充制造商" not in all_text


def test_renderer_shows_evidence_ids_for_each_relationship(tmp_path):
    payload = confirmed_payload()
    document = Document(render_report(payload, tmp_path / "out.docx"))
    text = document_text(document)
    assert "商标权利人：杭州星驰汽车用品有限公司（已确认；中；证据：EVD-001" in text


def test_renderer_requires_valid_style_tokens_and_applies_custom_tokens(tmp_path):
    style = json.loads(DEFAULT_STYLE_PATH.read_text(encoding="utf-8"))
    style["page"]["margin_left_mm"] = 25
    style["fonts"]["latin"] = "Custom Latin"
    style["sizes_pt"]["body"] = 11.5
    style["colors"]["primary"] = "123456"
    style_path = tmp_path / "style.json"
    style_path.write_text(json.dumps(style), encoding="utf-8")

    document = Document(render_report(confirmed_payload(), tmp_path / "out.docx", style_path))
    assert document.sections[0].left_margin.mm == pytest.approx(25, abs=0.1)
    normal = document.styles["Normal"]
    assert normal.font.name == "Custom Latin"
    assert normal.font.size.pt == pytest.approx(11.5)
    assert document.styles["Heading 1"].font.color.rgb == RGBColor.from_string("123456")
    assert document.styles["Normal"]._element.rPr.rFonts.get(qn("w:eastAsia")) == "Microsoft YaHei"

    with pytest.raises(ValueError, match="STYLE_READ_ERROR"):
        render_report(confirmed_payload(), tmp_path / "missing.docx", tmp_path / "missing.json")
    style_path.write_text("{not valid json", encoding="utf-8")
    with pytest.raises(ValueError, match="STYLE_JSON_ERROR"):
        render_report(confirmed_payload(), tmp_path / "invalid.docx", style_path)
    style_path.write_text(json.dumps({"page": {}}), encoding="utf-8")
    with pytest.raises(ValueError, match="STYLE_TOKEN_MISSING"):
        render_report(confirmed_payload(), tmp_path / "incomplete.docx", style_path)


@pytest.mark.parametrize(
    ("group", "key", "value"),
    [
        ("page", "margin_left_mm", True),
        ("page", "margin_right_mm", True),
        ("page", "margin_top_mm", True),
        ("page", "margin_bottom_mm", True),
        ("sizes_pt", "body", True),
        ("sizes_pt", "title", True),
        ("sizes_pt", "subtitle", True),
        ("sizes_pt", "heading1", True),
        ("sizes_pt", "heading2", True),
        ("sizes_pt", "table", True),
        ("spacing_pt", "body_after", True),
        ("spacing_pt", "heading1_before", True),
        ("spacing_pt", "heading1_after", True),
        ("spacing_pt", "heading2_before", True),
        ("spacing_pt", "heading2_after", True),
        ("table", "cell_margin_twips", True),
        ("page", "margin_left_mm", -1),
        ("page", "margin_left_mm", 101),
        ("sizes_pt", "body", 0),
        ("sizes_pt", "body", 73),
        ("spacing_pt", "body_after", 0),
        ("spacing_pt", "body_after", -1),
        ("spacing_pt", "body_after", 73),
        ("table", "cell_margin_twips", 0),
        ("table", "cell_margin_twips", -1),
        ("table", "cell_margin_twips", 1441),
    ],
)
def test_renderer_rejects_invalid_numeric_style_tokens(tmp_path, group, key, value):
    style = json.loads(DEFAULT_STYLE_PATH.read_text(encoding="utf-8"))
    style[group][key] = value
    style_path = tmp_path / "style.json"
    style_path.write_text(json.dumps(style), encoding="utf-8")

    with pytest.raises(ValueError, match="STYLE_TOKEN_INVALID"):
        render_report(confirmed_payload(), tmp_path / "out.docx", style_path)


def test_renderer_rejects_non_boolean_header_repeat(tmp_path):
    style = json.loads(DEFAULT_STYLE_PATH.read_text(encoding="utf-8"))
    style["table"]["header_repeat"] = 1
    style_path = tmp_path / "style.json"
    style_path.write_text(json.dumps(style), encoding="utf-8")

    with pytest.raises(ValueError, match="STYLE_TOKEN_INVALID table.header_repeat"):
        render_report(confirmed_payload(), tmp_path / "out.docx", style_path)


def test_renderer_sets_a4_fixed_table_geometry_repeating_headers_and_page_field(tmp_path):
    output = render_report(confirmed_payload(), tmp_path / "out.docx")
    document = Document(output)
    assert document.sections[0].page_width.mm == pytest.approx(210, abs=0.1)
    assert document.sections[0].page_height.mm == pytest.approx(297, abs=0.1)
    root = ET.fromstring(xml_part(output, "word/document.xml"))
    assert root.findall(".//w:tbl/w:tblPr/w:tblLayout[@w:type='fixed']", NS)
    assert root.findall(".//w:tr/w:trPr/w:tblHeader", NS)
    assert root.findall(".//w:instrText[.=' PAGE ']", NS) or xml_part(output, "word/footer1.xml")
    content_width = document.sections[0].page_width.twips - document.sections[0].left_margin.twips - document.sections[0].right_margin.twips
    for table in root.findall(".//w:tbl", NS):
        grid = [int(column.get(f"{{{W_NS}}}w")) for column in table.findall("w:tblGrid/w:gridCol", NS)]
        table_properties = table.find("w:tblPr", NS)
        width = int(table_properties.find("w:tblW", NS).get(f"{{{W_NS}}}w"))
        indent = int(table_properties.find("w:tblInd", NS).get(f"{{{W_NS}}}w"))
        assert indent + width <= content_width
        assert sum(grid) == width
        for row in table.findall("w:tr", NS):
            cells = row.findall("w:tc", NS)
            cell_widths = [int(cell.find("w:tcPr/w:tcW", NS).get(f"{{{W_NS}}}w")) for cell in cells]
            assert cell_widths == grid
            assert sum(cell_widths) == width


def test_renderer_is_deterministic_for_document_style_and_settings_xml(tmp_path):
    payload = confirmed_payload()
    first = render_report(payload, tmp_path / "first.docx")
    second = render_report(payload, tmp_path / "second.docx")
    for part in ("word/document.xml", "word/styles.xml", "word/settings.xml"):
        assert xml_part(first, part) == xml_part(second, part)


def test_cli_overwrites_only_requested_output_and_returns_one_on_invalid_json(tmp_path):
    input_path = tmp_path / "confirmed.json"
    output_path = tmp_path / "report.docx"
    sibling = tmp_path / "keep.txt"
    sibling.write_text("retain", encoding="utf-8")
    input_path.write_text(json.dumps(confirmed_payload(), ensure_ascii=False), encoding="utf-8")
    output_path.write_bytes(b"old")
    command = [sys.executable, "scripts/render_evidence_review_report.py", str(input_path), str(output_path)]
    result = subprocess.run(command, cwd=SKILL_ROOT, capture_output=True, text=True, check=False)
    assert result.returncode == 0
    assert output_path.exists() and output_path.read_bytes() != b"old"
    assert sibling.read_text(encoding="utf-8") == "retain"

    input_path.write_text(json.dumps(cloned_fixture("minimal-valid-review.json"), ensure_ascii=False), encoding="utf-8")
    before = output_path.read_bytes()
    result = subprocess.run(command, cwd=SKILL_ROOT, capture_output=True, text=True, check=False)
    assert result.returncode == 1
    assert "NOT_CONFIRMED" in result.stdout
    assert output_path.read_bytes() == before
