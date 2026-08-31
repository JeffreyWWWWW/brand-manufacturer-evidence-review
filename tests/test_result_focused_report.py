from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import sys

ROOT = Path(__file__).resolve().parents[1]
SKILL_ROOT = ROOT / "plugins" / "brand-manufacturer-evidence-review" / "skills" / "brand-manufacturer-evidence-review"
sys.path.insert(0, str(SKILL_ROOT))

from scripts.render_evidence_review_report import add_brand_sections, add_database_access, add_report_note


def test_report_note_keeps_metadata_but_database_access_is_not_report_content():
    document = Document()
    tokens = {
        "fonts": {"east_asia": "Microsoft YaHei", "latin": "Arial"},
        "sizes_pt": {"body": 10, "table": 8, "subtitle": 9},
        "colors": {"primary": "1F4E79", "text": "222222"},
    }
    payload = {
        "Skill版本": "1.1.0",
        "规范版本": "2026-08-01",
        "质量摘要": {"总体": {"品牌数量": 1, "冲突数": 0, "待补证数": 1}},
        "用户确认": {},
        "数据库访问记录": [{"数据库": "USPTO", "官方入口": "https://example.com", "访问时间": "2026-08-31T10:00:00+08:00", "访问结果": "已访问", "使用建议": "继续核验"}],
    }

    add_report_note(document, payload, tokens)

    add_database_access(document, payload, tokens)

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Skill版本：1.1.0" in text
    assert "数据库访问" not in text
    assert "访问时间" not in text


def test_brand_evidence_table_uses_result_focused_labels():
    document = Document()
    tokens = {
        "fonts": {"east_asia": "Microsoft YaHei", "latin": "Arial"},
        "sizes_pt": {"body": 10, "table": 8, "subtitle": 9, "heading1": 14, "heading2": 12},
        "colors": {"primary": "1F4E79", "text": "222222", "secondary": "D9EAF7", "border": "B7C9D6", "muted": "666666"},
        "table": {"cell_margin_twips": 120, "header_repeat": True},
    }
    payload = {
        "调查范围": {"目标品牌": [{"原始名称": "Demo", "规范名称": "Demo"}]},
        "品牌复核结果": [{
            "品牌": {"原始名称": "Demo", "规范名称": "Demo"},
            "查询结果摘要": "Demo 的运营主体已识别。",
            "主体关系": {},
            "制造关系": {"制造模式": "未知", "SKU适用限制": []},
            "结论评价": {"可靠性等级": "中", "可靠性依据": "证据有限", "人工复核建议": "补充资料"},
            "关键说明": [],
            "主要来源": [{
                "证据编号": "EVD-001", "来源名称": "Demo official", "URL": "https://example.com", "来源类别": "品牌官方页面",
                "访问日期": "2026-08-31", "支持结论": ["运营主体"], "证据等级": "官方企业资料", "原文摘录": "Demo", "页面定位": "About"
            }],
        }],
        "主体索引": [],
    }

    add_brand_sections(document, payload, tokens)

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    for forbidden in ("实际查询路径", "检索词", "失败原因", "数据库访问", "访问时间"):
        assert forbidden not in text
    assert "证据日期" in text


def test_brand_evidence_urls_are_clickable_hyperlinks():
    document = Document()
    tokens = {
        "fonts": {"east_asia": "Microsoft YaHei", "latin": "Arial"},
        "sizes_pt": {"body": 10, "table": 8, "subtitle": 9, "heading1": 14, "heading2": 12},
        "colors": {"primary": "1F4E79", "text": "222222", "secondary": "D9EAF7", "border": "B7C9D6", "muted": "666666"},
        "table": {"cell_margin_twips": 120, "header_repeat": True},
    }
    url = "https://example.com/evidence?id=1"
    payload = {
        "调查范围": {"目标品牌": [{"原始名称": "Demo", "规范名称": "Demo"}]},
        "品牌复核结果": [{
            "品牌": {"原始名称": "Demo", "规范名称": "Demo"},
            "查询结果摘要": "Demo 的运营主体已识别。",
            "主体关系": {},
            "制造关系": {"制造模式": "未知", "SKU适用限制": []},
            "结论评价": {"可靠性等级": "中", "可靠性依据": "证据有限", "人工复核建议": "补充资料"},
            "关键说明": [],
            "主要来源": [{
                "证据编号": "EVD-001", "来源名称": "Demo official", "URL": url, "来源类别": "品牌官方页面",
                "访问日期": "2026-08-31", "支持结论": ["运营主体"], "证据等级": "官方企业资料", "原文摘录": "Demo", "页面定位": "About"
            }, {
                "证据编号": "EVD-002", "来源名称": "Offline filing", "URL": None, "来源类别": "离线文件",
                "访问日期": "2026-08-31", "支持结论": ["运营主体"], "证据等级": "官方登记", "原文摘录": "Demo", "页面定位": "第 1 页"
            }],
        }],
        "主体索引": [],
    }

    add_brand_sections(document, payload, tokens)

    url_cell = document.tables[0].cell(1, 2)
    hyperlinks = url_cell._tc.findall(".//" + qn("w:hyperlink"))
    assert len(hyperlinks) == 1
    relationship_id = hyperlinks[0].get(qn("r:id"))
    relationship = document.part.rels[relationship_id]
    assert relationship.is_external
    assert relationship.target_ref == url
    assert url_cell.text == url
    run_properties = hyperlinks[0].find(".//" + qn("w:rPr"))
    assert run_properties.find(qn("w:sz")).get(qn("w:val")) == "16"
    assert run_properties.find(qn("w:rFonts")).get(qn("w:eastAsia")) == "Microsoft YaHei"
    empty_url_cell = document.tables[0].cell(2, 2)
    assert empty_url_cell.text == ""
    assert empty_url_cell._tc.findall(".//" + qn("w:hyperlink")) == []
